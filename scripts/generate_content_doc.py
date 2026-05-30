"""
VIABIL Website Content Extractor
Generates an organized Word document with all site written content,
mapped to their source file locations for easy editing and re-import.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─────────────────────────────────────────────
#  Helpers
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_page_break(doc):
    doc.add_page_break()

def add_source_note(doc, source):
    p = doc.add_paragraph()
    run = p.add_run(f"Source: {source}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.italic = True
    p.paragraph_format.space_after = Pt(4)

def add_field_row(table, label, content, source=""):
    row = table.add_row()
    label_cell = row.cells[0]
    content_cell = row.cells[1]
    source_cell = row.cells[2]

    label_p = label_cell.paragraphs[0]
    label_run = label_p.add_run(label)
    label_run.bold = True
    label_run.font.size = Pt(9)
    label_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    set_cell_bg(label_cell, "F0F4FF")

    content_p = content_cell.paragraphs[0]
    content_run = content_p.add_run(content)
    content_run.font.size = Pt(10)

    source_p = source_cell.paragraphs[0]
    source_run = source_p.add_run(source)
    source_run.font.size = Pt(8)
    source_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    source_run.font.italic = True

def make_section_table(doc):
    table = doc.add_table(rows=0, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Set column widths
    for i, width in enumerate([Cm(4), Cm(10), Cm(4)]):
        for cell in table.columns[i].cells:
            cell.width = width
    # Header row
    header = table.add_row()
    for i, text in enumerate(["Field", "Content", "Source"]):
        cell = header.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, "1A1A2E")
    return table

def add_page_heading(doc, title, file_path):
    p = doc.add_heading(title, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    add_source_note(doc, file_path)

def add_section_heading(doc, title):
    p = doc.add_heading(title, level=2)
    if p.runs:
        p.runs[0].font.color.rgb = RGBColor(0x2D, 0x5A, 0xFF)

def add_subsection_heading(doc, title):
    p = doc.add_heading(title, level=3)
    if p.runs:
        p.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x66)

def spacer(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

# ─────────────────────────────────────────────
#  Document setup
# ─────────────────────────────────────────────

doc = Document()

# Margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Title page
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("VIABIL Website")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub_p.add_run("Written Content Review Document")
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(0x2D, 0x5A, 0xFF)

doc.add_paragraph()
note_p = doc.add_paragraph()
note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
note_run = note_p.add_run(
    "This document contains all written content from the VIABIL website.\n"
    "Edit the Content column freely. The Source column shows where each text\n"
    "lives in the codebase so changes can be applied precisely."
)
note_run.font.size = Pt(10)
note_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
note_run.font.italic = True

doc.add_paragraph()
instructions = doc.add_paragraph()
instructions.alignment = WD_ALIGN_PARAGRAPH.CENTER
inst_run = instructions.add_run(
    "Column guide: Field = text identifier  |  Content = editable text  |  Source = file:line"
)
inst_run.font.size = Pt(9)
inst_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
#  1. GLOBAL / METADATA
# ─────────────────────────────────────────────────────────────────────────────
add_page_heading(doc, "1. Global — Metadata & SEO", "app/layout.tsx")
t = make_section_table(doc)
add_field_row(t, "Page title",         "VIABIL | Inteligência financeira para o mercado imobiliário", "layout.tsx:14")
add_field_row(t, "Meta description",   "O VIABIL é a plataforma de referência em viabilidade econômico-financeira para incorporadoras, loteadoras e desenvolvedores imobiliários. Do terreno ao resultado.", "layout.tsx:15")
add_field_row(t, "OG title",           "VIABIL | Inteligência financeira para o mercado imobiliário", "layout.tsx:19")
add_field_row(t, "OG description",     "Do terreno ao resultado: decisões mais seguras para incorporadoras, loteadoras e desenvolvedores imobiliários.", "layout.tsx:20")
add_field_row(t, "OG site name",       "VIABIL", "layout.tsx:21")
add_field_row(t, "Twitter title",      "VIABIL | Inteligência financeira para o mercado imobiliário", "layout.tsx:25")
add_field_row(t, "Twitter description","Do terreno ao resultado: decisões mais seguras para incorporadoras, loteadoras e desenvolvedores imobiliários.", "layout.tsx:26")

spacer(doc)

# ─────────────────────────────────────────────────────────────────────────────
#  2. NAVBAR
# ─────────────────────────────────────────────────────────────────────────────
add_page_heading(doc, "2. Navigation (Navbar)", "components/marketing/Navbar.tsx")

add_section_heading(doc, "Nav links")
t = make_section_table(doc)
add_field_row(t, "Nav link 1", "Plataforma",  "Navbar.tsx:8")
add_field_row(t, "Nav link 2", "Módulos",     "Navbar.tsx:9")
add_field_row(t, "Nav link 3", "Segmentos",   "Navbar.tsx:10")
add_field_row(t, "Nav link 4", "Versões",     "Navbar.tsx:11")
add_field_row(t, "Nav link 5", "Serviços",    "Navbar.tsx:12")
add_field_row(t, "Nav link 6", "Sobre",       "Navbar.tsx:13")
spacer(doc)

add_section_heading(doc, "CTA button")
t = make_section_table(doc)
add_field_row(t, "CTA button", "Solicitar demonstração", "Navbar.tsx:48")
spacer(doc)

# ─────────────────────────────────────────────────────────────────────────────
#  3. FOOTER
# ─────────────────────────────────────────────────────────────────────────────
add_page_heading(doc, "3. Footer", "components/marketing/Footer.tsx")
t = make_section_table(doc)
add_field_row(t, "Tagline",        "Software de viabilidade econômico-financeira para empreendimentos imobiliários. Conhecimento e tecnologia para decisões com segurança.", "Footer.tsx:31")
add_field_row(t, "Column 1 title", "Produto",                              "Footer.tsx:38")
add_field_row(t, "Column 1 link 1","Plataforma",                           "Footer.tsx:39")
add_field_row(t, "Column 1 link 2","Módulos",                              "Footer.tsx:40")
add_field_row(t, "Column 1 link 3","Versões",                              "Footer.tsx:41")
add_field_row(t, "Column 1 link 4","Serviços",                             "Footer.tsx:42")
add_field_row(t, "Column 2 title", "Mercado",                              "Footer.tsx:45")
add_field_row(t, "Column 2 link 1","Segmentos",                            "Footer.tsx:46")
add_field_row(t, "Column 2 link 2","BDK Solutions",                        "Footer.tsx:47")
add_field_row(t, "Column 2 link 3","Contato",                              "Footer.tsx:48")
add_field_row(t, "Copyright",      "© [year] BDK Solutions. Todos os direitos reservados.", "Footer.tsx:54")
add_field_row(t, "Trademark",      "VIABIL® é uma marca da BDK Solutions.", "Footer.tsx:55")
spacer(doc)

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
#  4. HOMEPAGE (landing-page.tsx)
# ─────────────────────────────────────────────────────────────────────────────
add_page_heading(doc, "4. Homepage", "components/landing-page.tsx")

add_section_heading(doc, "Hero section")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",    "Plataforma",                                                         "landing-page.tsx — hero eyebrow")
add_field_row(t, "Headline",   "Inteligência financeira para todo o ciclo imobiliário.",              "landing-page.tsx — hero h1")
add_field_row(t, "Subheadline","O VIABIL nasceu dentro da incorporação imobiliária para simular cenários, analisar indicadores e acompanhar a rentabilidade do negócio do Go/No-Go ao previsto x realizado.", "landing-page.tsx — hero subtitle")
add_field_row(t, "CTA primary","Solicitar demonstração",                                             "landing-page.tsx — hero CTA 1")
add_field_row(t, "CTA secondary","Ver módulos",                                                      "landing-page.tsx — hero CTA 2")
spacer(doc)

add_section_heading(doc, "Social proof / clients strip")
t = make_section_table(doc)
add_field_row(t, "Label",      "Empresas que confiam no VIABIL",    "landing-page.tsx — clients label")
spacer(doc)

add_section_heading(doc, "Problem / positioning section")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",    "O que é",                           "landing-page.tsx — o-que-e eyebrow")
add_field_row(t, "Heading",    "Uma referência de mercado, não uma planilha mais bonita.", "landing-page.tsx — o-que-e heading")
add_field_row(t, "Subheading", "O VIABIL concentra conhecimento prático do setor, modelos financeiros testados e governança para decisões de alto impacto em incorporação residencial, casas, loteamentos e outros segmentos.", "landing-page.tsx — o-que-e subheading")
add_field_row(t, "Pill 1",     "DNA imobiliário",                   "landing-page.tsx — pill 1")
add_field_row(t, "Pill 2",     "Padrão entre parceiros",            "landing-page.tsx — pill 2")
add_field_row(t, "Pill 3",     "Cenários vivos",                    "landing-page.tsx — pill 3")
add_field_row(t, "Pill 4",     "Governança de premissas",           "landing-page.tsx — pill 4")
spacer(doc)

add_section_heading(doc, "Ciclo VIABIL section")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",    "Ciclo VIABIL",                      "landing-page.tsx — ciclo eyebrow")
add_field_row(t, "Heading",    "Do terreno ao resultado, com a mesma visão gerencial.", "landing-page.tsx — ciclo heading")
add_field_row(t, "Subheading", "A análise não para na aprovação. O ciclo continua quando as premissas mudam, a obra avança e a empresa precisa corrigir rota.", "landing-page.tsx — ciclo subheading")
add_field_row(t, "Step 1 label","Captação",                          "CicloOrbital.tsx:6")
add_field_row(t, "Step 1 text", "Organize oportunidades, documentos, mapas e histórico de negociação antes da decisão de compra.", "CicloOrbital.tsx:7")
add_field_row(t, "Step 2 label","Viabilidade",                       "CicloOrbital.tsx:14")
add_field_row(t, "Step 2 text", "Modele VGV, custos, financiamento, permutas, velocidade de vendas, indicadores e cenários.", "CicloOrbital.tsx:15")
add_field_row(t, "Step 3 label","Decisão",                           "CicloOrbital.tsx:22")
add_field_row(t, "Step 3 text", "Leve relatórios consistentes para sócios, investidores, comitês e conselhos.", "CicloOrbital.tsx:23")
add_field_row(t, "Step 4 label","Acompanhamento",                    "CicloOrbital.tsx:30")
add_field_row(t, "Step 4 text", "Compare planejado, revisado e realizado para agir antes que o resultado se perca.", "CicloOrbital.tsx:31")
add_field_row(t, "Step 5 label","Replanejamento",                    "CicloOrbital.tsx:38")
add_field_row(t, "Step 5 text", "A partir dos resultados do acompanhamento, simule ajustes de premissas para recuperar ou superar as metas do empreendimento.", "CicloOrbital.tsx:39")
spacer(doc)

add_section_heading(doc, "Depoimentos (Testimonials)")
t = make_section_table(doc)
add_field_row(t, "Section eyebrow", "Depoimentos",                  "DepoimentosSection.tsx:116")
add_field_row(t, "Section heading", "Como clientes descrevem o impacto do VIABIL.", "DepoimentosSection.tsx:117")
spacer(doc)
doc.add_paragraph("— Testimonial 1")
t = make_section_table(doc)
add_field_row(t, "Quote",    "O VIABIL é sem sombra de dúvidas o principal software de viabilidade de empreendimentos imobiliários do Brasil. É um instrumento importante para as empresas que pretendem melhorar a governança.", "DepoimentosSection.tsx:8")
add_field_row(t, "Name",     "Felipe Cavalcante",    "DepoimentosSection.tsx:13")
add_field_row(t, "Role",     "Presidente · ADITBrasil", "DepoimentosSection.tsx:14")
spacer(doc)
doc.add_paragraph("— Testimonial 2")
t = make_section_table(doc)
add_field_row(t, "Quote",    "Nestes últimos anos, o mercado imobiliário sofreu uma retração, mas não parou... Diante deste novo cenário, o uso do VIABIL tornou-se ainda mais estratégico e tem papel decisivo na hora de decidirmos em que negócio colocaremos nossas energias e nossos recursos.", "DepoimentosSection.tsx:17")
add_field_row(t, "Name",     "Hélio Abreu",           "DepoimentosSection.tsx:22")
add_field_row(t, "Role",     "Sócio-Diretor · Record","DepoimentosSection.tsx:23")
spacer(doc)
doc.add_paragraph("— Testimonial 3")
t = make_section_table(doc)
add_field_row(t, "Quote",    "O VIABIL é uma ferramenta indispensável no dia a dia da empresa, proporcionando controle, organização, segurança e agilidade para os negócios.", "DepoimentosSection.tsx:26")
add_field_row(t, "Name",     "Camilo Vieira dos Santos", "DepoimentosSection.tsx:31")
add_field_row(t, "Role",     "HM Engenharia",            "DepoimentosSection.tsx:32")
spacer(doc)
doc.add_paragraph("— Testimonial 4")
t = make_section_table(doc)
add_field_row(t, "Quote",    "Posso dizer com certeza e garantia que o sistema VIABIL é uma ferramenta fundamental e imprescindível na análise de novos negócios para que possamos aferir, definir, planejar e conseguir alcançar nossos almejados resultados em cada empreendimento.", "DepoimentosSection.tsx:35")
add_field_row(t, "Name",     "Wilson Sequeira",               "DepoimentosSection.tsx:40")
add_field_row(t, "Role",     "Diretor de Incorporação · INTERRIO", "DepoimentosSection.tsx:41")
spacer(doc)
doc.add_paragraph("— Testimonial 5")
t = make_section_table(doc)
add_field_row(t, "Quote",    "Nesses 15 anos de mercado, tive a oportunidade de poder trabalhar com este sistema em todas as empresas onde passei. O VIABIL acompanhou as mudanças, desenvolveu novas ferramentas e colaborou com o crescimento do Real Estate em todo o Brasil.", "DepoimentosSection.tsx:44")
add_field_row(t, "Name",     "Greco G Montagna",              "DepoimentosSection.tsx:49")
add_field_row(t, "Role",     "Gerente Comercial Real Estate · BTGPactual", "DepoimentosSection.tsx:50")
spacer(doc)
doc.add_paragraph("— Testimonial 6")
t = make_section_table(doc)
add_field_row(t, "Quote",    "Seus recursos atendem às necessidades da empresa, padronizando procedimentos, análises físico-financeiras e acompanhamento de resultados.", "DepoimentosSection.tsx:53")
add_field_row(t, "Name",     "Equipe de Planejamento",        "DepoimentosSection.tsx:58")
add_field_row(t, "Role",     "Cyrela",                        "DepoimentosSection.tsx:59")
spacer(doc)
doc.add_paragraph("— Testimonial 7")
t = make_section_table(doc)
add_field_row(t, "Quote",    "O VIABIL é um aliado da empresa, dando agilidade ao processo e fornecendo informações claras e objetivas que permitem aos nossos diretores tomar decisões mais seguras com relação aos nossos investimentos.", "DepoimentosSection.tsx:62")
add_field_row(t, "Name",     "Diretoria de Investimentos",    "DepoimentosSection.tsx:67")
add_field_row(t, "Role",     "Rodobens Negócios Imobiliários","DepoimentosSection.tsx:68")
spacer(doc)
doc.add_paragraph("— Testimonial 8")
t = make_section_table(doc)
add_field_row(t, "Quote",    "Com o VIABIL, conseguimos parametrizar nossos estudos, aumentar nossa assertividade e controlar o acesso a múltiplos usuários, sem perder a confiabilidade nos resultados e ainda conseguimos trocar informações com rapidez e transparência com outras incorporadoras.", "DepoimentosSection.tsx:71")
add_field_row(t, "Name",     "Novos Negócios",                "DepoimentosSection.tsx:76")
add_field_row(t, "Role",     "Porto Ferraz Construtora",      "DepoimentosSection.tsx:77")
spacer(doc)
doc.add_paragraph("— Testimonial 9")
t = make_section_table(doc)
add_field_row(t, "Quote",    "O VIABIL é uma ferramenta indispensável no dia a dia de nossa empresa, seja para cadastrar terrenos, padronizar os estudos de viabilidade econômica, tomadas de decisões de investimento e controle dos nossos resultados.", "DepoimentosSection.tsx:80")
add_field_row(t, "Name",     "Equipe Técnica",                "DepoimentosSection.tsx:85")
add_field_row(t, "Role",     "Cury Construtora",              "DepoimentosSection.tsx:86")
spacer(doc)
doc.add_paragraph("— Testimonial 10")
t = make_section_table(doc)
add_field_row(t, "Quote",    "A ferramenta acompanha mudanças do mercado e proporciona respostas rápidas sem perder a capacidade de analisar diversas variáveis.", "DepoimentosSection.tsx:89")
add_field_row(t, "Name",     "Real Estate",                   "DepoimentosSection.tsx:94")
add_field_row(t, "Role",     "BTG Pactual",                   "DepoimentosSection.tsx:95")
spacer(doc)
doc.add_paragraph("— Testimonial 11")
t = make_section_table(doc)
add_field_row(t, "Quote",    "O treinamento foi flexibilizado às nossas necessidades de horário e conteúdo, sempre com atenção à nossa realidade.", "DepoimentosSection.tsx:98")
add_field_row(t, "Name",     "Treinamento e Implantação",     "DepoimentosSection.tsx:103")
add_field_row(t, "Role",     "Sequóia",                       "DepoimentosSection.tsx:104")
spacer(doc)
doc.add_paragraph("— Testimonial 12")
t = make_section_table(doc)
add_field_row(t, "Quote",    "An outstanding and dynamic training, that I'll definitely recommend to all my partners. It's extremely important because VIABIL covers deep content that we must understand to operate better and efficiently.", "DepoimentosSection.tsx:107")
add_field_row(t, "Name",     "Equipe",                        "DepoimentosSection.tsx:112")
add_field_row(t, "Role",     "Paroma Incorporações",          "DepoimentosSection.tsx:113")
spacer(doc)

add_section_heading(doc, "CTA Band (Homepage)")
t = make_section_table(doc)
add_field_row(t, "Heading",    "Pronto para conhecer o VIABIL?",                                 "CTABand.tsx:9")
add_field_row(t, "Subheading", "Solicite uma demonstração e veja como o VIABIL apoia decisões financeiras críticas, do terreno ao resultado.", "CTABand.tsx:10")
add_field_row(t, "Button 1",   "Solicitar demonstração",                                         "CTABand.tsx:18")
add_field_row(t, "Button 2",   "Falar com especialista",                                         "CTABand.tsx:23")
spacer(doc)

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
#  5. PLATAFORMA PAGE
# ─────────────────────────────────────────────────────────────────────────────
add_page_heading(doc, "5. Plataforma Page", "app/(public)/plataforma/page.tsx")

add_section_heading(doc, "Page metadata")
t = make_section_table(doc)
add_field_row(t, "Page title",       "Plataforma | VIABIL",                                          "plataforma/page.tsx:8")
add_field_row(t, "Meta description", "Conheça o VIABIL: a referência em viabilidade econômico-financeira para o ciclo completo do empreendimento imobiliário.", "plataforma/page.tsx:9")
spacer(doc)

add_section_heading(doc, "Hero")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",     "Plataforma",                                               "plataforma/page.tsx:73")
add_field_row(t, "H1",         "Inteligência financeira para todo o ciclo imobiliário.",    "plataforma/page.tsx:75")
add_field_row(t, "Subtitle",   "O VIABIL nasceu dentro da incorporação imobiliária para simular cenários, analisar indicadores e acompanhar a rentabilidade do negócio do Go/No-Go ao previsto x realizado.", "plataforma/page.tsx:77")
add_field_row(t, "CTA 1",      "Solicitar demonstração",                                   "plataforma/page.tsx:82")
add_field_row(t, "CTA 2",      "Ver módulos",                                              "plataforma/page.tsx:85")
spacer(doc)

add_section_heading(doc, "O que é")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",    "O que é",                                                   "plataforma/page.tsx:98")
add_field_row(t, "Heading",    "Uma referência de mercado, não uma planilha mais bonita.",  "plataforma/page.tsx:100")
add_field_row(t, "Subheading", "O VIABIL concentra conhecimento prático do setor, modelos financeiros testados e governança para decisões de alto impacto em incorporação residencial, casas, loteamentos e outros segmentos.", "plataforma/page.tsx:102")
add_field_row(t, "Pill 1",     "DNA imobiliário",                                           "plataforma/page.tsx:105")
add_field_row(t, "Pill 2",     "Padrão entre parceiros",                                   "plataforma/page.tsx:106")
add_field_row(t, "Pill 3",     "Cenários vivos",                                           "plataforma/page.tsx:107")
add_field_row(t, "Pill 4",     "Governança de premissas",                                  "plataforma/page.tsx:108")
spacer(doc)

add_section_heading(doc, "Ciclo VIABIL")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",    "Ciclo VIABIL",                                              "plataforma/page.tsx:116")
add_field_row(t, "Heading",    "Do terreno ao resultado, com a mesma visão gerencial.",     "plataforma/page.tsx:118")
add_field_row(t, "Subheading", "A análise não para na aprovação. O ciclo continua quando as premissas mudam, a obra avança e a empresa precisa corrigir rota.", "plataforma/page.tsx:120")
spacer(doc)

add_section_heading(doc, "Planilhas x VIABIL")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",           "Planilhas x VIABIL",                                        "plataforma/page.tsx:135")
add_field_row(t, "Heading",           "A discussão sai da fórmula e volta para a decisão.",         "plataforma/page.tsx:137")
add_field_row(t, "Subheading",        "Em negócios imobiliários, a fragilidade está no controle de premissas, versões, indicadores e histórico. O VIABIL reduz esse ruído.", "plataforma/page.tsx:139")
add_field_row(t, "Left col heading",  "Quando a análise fica em planilhas",                         "plataforma/page.tsx:144")
add_field_row(t, "Left Governança",   "Arquivos circulam sem rastreabilidade clara.",                "plataforma/page.tsx:146")
add_field_row(t, "Left Método",       "Cada equipe pode calcular de um jeito.",                     "plataforma/page.tsx:148")
add_field_row(t, "Left Cenários",     "Testar mudanças exige refazer muito trabalho.",               "plataforma/page.tsx:150")
add_field_row(t, "Left Portfólio",    "A visão consolidada depende de recortes manuais.",            "plataforma/page.tsx:152")
add_field_row(t, "Right col heading", "Quando a análise roda no VIABIL",                            "plataforma/page.tsx:155")
add_field_row(t, "Right Governança",  "Premissas, versões e usuários ficam organizados.",            "plataforma/page.tsx:157")
add_field_row(t, "Right Método",      "A empresa trabalha com cálculo testado e padrão de mercado.", "plataforma/page.tsx:159")
add_field_row(t, "Right Cenários",    "Stress-cenários e simulações mantêm a decisão viva.",         "plataforma/page.tsx:161")
add_field_row(t, "Right Portfólio",   "Projetos e oportunidades alimentam uma visão executiva.",     "plataforma/page.tsx:163")
spacer(doc)

add_section_heading(doc, "Pilares")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",    "Pilares",                                                    "plataforma/page.tsx:175")
add_field_row(t, "Heading",    "Valor, flexibilidade, parametrização e confiança.",          "plataforma/page.tsx:177")
add_field_row(t, "Subheading", "Esses quatro princípios precisam aparecer em toda a experiência porque explicam por que o VIABIL é diferente de soluções genéricas.", "plataforma/page.tsx:179")
add_field_row(t, "Pilar 1 title","Valor agregado",                                           "plataforma/page.tsx:183")
add_field_row(t, "Pilar 1 text","Cada estudo passa a sustentar uma decisão de negócio, não apenas um cálculo isolado.", "plataforma/page.tsx:184")
add_field_row(t, "Pilar 2 title","Flexibilidade",                                            "plataforma/page.tsx:186")
add_field_row(t, "Pilar 2 text","A plataforma se adapta a segmentos, estruturas societárias, modelos financeiros e regiões.", "plataforma/page.tsx:187")
add_field_row(t, "Pilar 3 title","Parametrização",                                           "plataforma/page.tsx:189")
add_field_row(t, "Pilar 3 text","Premissas, indicadores, relatórios e modelos seguem a forma de trabalho da empresa.", "plataforma/page.tsx:190")
add_field_row(t, "Pilar 4 title","Confiança",                                                "plataforma/page.tsx:192")
add_field_row(t, "Pilar 4 text","Décadas de uso no mercado reduzem discussões sobre fórmulas e elevam o debate sobre premissas.", "plataforma/page.tsx:193")
spacer(doc)

add_section_heading(doc, "Decisão Contínua")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",    "Decisão contínua",                                          "plataforma/page.tsx:194")
add_field_row(t, "Heading",    "As decisões não são mais estáticas.",                        "plataforma/page.tsx:196")
add_field_row(t, "Subheading", "O VIABIL ajuda a empresa a avaliar a saúde dos empreendimentos, a necessidade de funding e o retorno esperado conforme o projeto evolui.", "plataforma/page.tsx:198")
add_field_row(t, "Item 1 title","Do ativo ao portfólio",                                     "plataforma/page.tsx:202")
add_field_row(t, "Item 1 text","O estudo individual precisa conversar com a necessidade de caixa, retorno esperado e exposição dos acionistas no tempo.", "plataforma/page.tsx:203")
add_field_row(t, "Item 2 title","Premissas vivas",                                           "plataforma/page.tsx:205")
add_field_row(t, "Item 2 text","Preço, custo, velocidade de vendas, financiamento e permutas mudam. A decisão precisa ser recalculável sem perder histórico.", "plataforma/page.tsx:206")
add_field_row(t, "Item 3 title","Discussão executiva",                                       "plataforma/page.tsx:208")
add_field_row(t, "Item 3 text","Relatórios e indicadores padronizados reduzem ruído técnico e ajudam comitês a discutir risco, retorno e alternativa de ação.", "plataforma/page.tsx:209")
spacer(doc)

add_section_heading(doc, "Análise de Sensibilidade")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",    "Análise de sensibilidade",                                  "plataforma/page.tsx:214")
add_field_row(t, "Heading",    "Como preço, custo e velocidade de vendas mudam o resultado.","plataforma/page.tsx:216")
add_field_row(t, "Subheading", "O VIABIL gera automaticamente o mapa de sensibilidade do estudo: uma matriz que mostra como os indicadores reagem a variações nas premissas críticas. Exemplo real abaixo, gerado pela plataforma.", "plataforma/page.tsx:218")
add_field_row(t, "Table note",  "Exemplo gerado pelo VIABIL para incorporação residencial SFH. VGV R$ 24,8M, 40 unidades, 24 meses de obra.", "plataforma/page.tsx:262")
add_field_row(t, "Link text",   "Abrir exemplo completo ↗",                                  "plataforma/page.tsx:265")
spacer(doc)

add_section_heading(doc, "CTA Band")
t = make_section_table(doc)
add_field_row(t, "Heading",    "Veja o ciclo completo em uma demonstração.",                 "plataforma/page.tsx:280")
add_field_row(t, "Subheading", "A melhor conversa começa com o seu tipo de empreendimento, suas premissas e as decisões que sua equipe precisa sustentar.", "plataforma/page.tsx:282")
spacer(doc)

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
#  6. MÓDULOS PAGE
# ─────────────────────────────────────────────────────────────────────────────
add_page_heading(doc, "6. Módulos Page", "app/(public)/modulos/page.tsx")

add_section_heading(doc, "Page metadata")
t = make_section_table(doc)
add_field_row(t, "Page title",       "Módulos | VIABIL",                                            "modulos/page.tsx:8")
add_field_row(t, "Meta description", "Conheça os módulos do VIABIL: Gestão de Terrenos, Viabilidade, Acompanhamento, Consolidação e Workflow de Tarefas.", "modulos/page.tsx:9")
spacer(doc)

add_section_heading(doc, "Hero")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",   "Módulos",                                                    "modulos/page.tsx:74")
add_field_row(t, "H1",        "Cinco módulos para uma visão contínua do negócio.",          "modulos/page.tsx:76")
add_field_row(t, "Subtitle",  "O VIABIL conecta originação, viabilidade, decisão, acompanhamento, consolidação e processo. A empresa deixa de analisar eventos isolados e passa a gerir o ciclo financeiro completo.", "modulos/page.tsx:78")
spacer(doc)

add_section_heading(doc, "Módulo 01 — Gestão de Terrenos")
t = make_section_table(doc)
add_field_row(t, "Tag",         "Originação e landbank",                                     "modulos/page.tsx module 1")
add_field_row(t, "Headline",    "A oportunidade entra organizada antes de virar estudo.",    "modulos/page.tsx module 1")
add_field_row(t, "Description", "Centraliza terrenos oferecidos e prospectados, documentos, imagens, dados urbanísticos, histórico de negociação, filtros e tarefas da equipe de Novos Negócios.", "modulos/page.tsx module 1")
add_field_row(t, "Fact 1",      "40+ filtros combinados",                                    "modulos/page.tsx module 1")
add_field_row(t, "Fact 2",      "Google Maps e documentos",                                  "modulos/page.tsx module 1")
add_field_row(t, "Fact 3",      "Histórico de negociação",                                   "modulos/page.tsx module 1")
add_field_row(t, "Fact 4",      "Link direto para Viabilidade",                              "modulos/page.tsx module 1")
spacer(doc)

add_section_heading(doc, "Módulo 02 — Viabilidade")
t = make_section_table(doc)
add_field_row(t, "Tag",         "Simulação financeira",                                      "modulos/page.tsx module 2")
add_field_row(t, "Headline",    "O motor principal para decisões de Go/No-Go.",              "modulos/page.tsx module 2")
add_field_row(t, "Description", "Projeta fluxo de caixa, indicadores e premissas para incorporação residencial, casas, loteamentos, MCMV, corporativo, logística, shopping e projetos mistos.", "modulos/page.tsx module 2")
add_field_row(t, "Fact 1",      "VGV, margem, VPL, TIR, MTIR e ROI",                        "modulos/page.tsx module 2")
add_field_row(t, "Fact 2",      "Stress-cenários em variáveis críticas",                     "modulos/page.tsx module 2")
add_field_row(t, "Fact 3",      "Premissas e modelos parametrizáveis",                       "modulos/page.tsx module 2")
add_field_row(t, "Fact 4",      "Relatórios exportáveis para Excel",                         "modulos/page.tsx module 2")
spacer(doc)

add_section_heading(doc, "Módulo 03 — Acompanhamento")
t = make_section_table(doc)
add_field_row(t, "Tag",         "Previsto x realizado",                                      "modulos/page.tsx module 3")
add_field_row(t, "Headline",    "Não basta acompanhar. Precisa agir.",                       "modulos/page.tsx module 3")
add_field_row(t, "Description", "Compara planejado, revisado e realizado, importa dados de ERPs ou planilhas e permite replanejar ações para buscar as metas definidas no estudo.", "modulos/page.tsx module 3")
add_field_row(t, "Fact 1",      "Previsto x revisado x realizado",                           "modulos/page.tsx module 3")
add_field_row(t, "Fact 2",      "Alertas de divergência",                                    "modulos/page.tsx module 3")
add_field_row(t, "Fact 3",      "Wizard de reprojeção",                                      "modulos/page.tsx module 3")
add_field_row(t, "Fact 4",      "Visão para sócios e investidores",                          "modulos/page.tsx module 3")
spacer(doc)

add_section_heading(doc, "Módulo 04 — Consolidação de Resultados")
t = make_section_table(doc)
add_field_row(t, "Tag",         "Portfólio",                                                  "modulos/page.tsx module 4")
add_field_row(t, "Headline",    "A visão executiva entre projetos, oportunidades e capital.", "modulos/page.tsx module 4")
add_field_row(t, "Description", "Consolida fluxos e indicadores de projetos em prospecção, desenvolvimento e modelos futuros para apoiar planejamento estratégico e decisões de alocação.", "modulos/page.tsx module 4")
add_field_row(t, "Fact 1",      "Fluxo consolidado",                                         "modulos/page.tsx module 4")
add_field_row(t, "Fact 2",      "Comparativo entre cenários",                                "modulos/page.tsx module 4")
add_field_row(t, "Fact 3",      "Necessidade de aporte no tempo",                            "modulos/page.tsx module 4")
add_field_row(t, "Fact 4",      "Ranking de oportunidades",                                  "modulos/page.tsx module 4")
spacer(doc)

add_section_heading(doc, "Módulo 05 — Workflow de Tarefas")
t = make_section_table(doc)
add_field_row(t, "Tag",         "Processo e governança",                                     "modulos/page.tsx module 5")
add_field_row(t, "Headline",    "Cada etapa com responsável, prazo e histórico.",             "modulos/page.tsx module 5")
add_field_row(t, "Description", "Gerencia atividades desde captação do terreno até chaves e recebíveis, com checklists, pendências por usuário e acompanhamento gerencial.", "modulos/page.tsx module 5")
add_field_row(t, "Fact 1",      "Etapas e responsáveis",                                     "modulos/page.tsx module 5")
add_field_row(t, "Fact 2",      "Pendências por usuário",                                    "modulos/page.tsx module 5")
add_field_row(t, "Fact 3",      "Lembretes por e-mail",                                      "modulos/page.tsx module 5")
add_field_row(t, "Fact 4",      "Histórico por terreno ou projeto",                          "modulos/page.tsx module 5")
spacer(doc)

add_section_heading(doc, "Integração section")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",    "Integração",                                                 "modulos/page.tsx:119")
add_field_row(t, "Heading",    "O ganho aparece quando os módulos conversam.",               "modulos/page.tsx:121")
add_field_row(t, "Subheading", "Uma oportunidade cadastrada em Terrenos vira estudo em Viabilidade. O estudo aprovado vira referência para Acompanhamento. Os resultados alimentam a Consolidação e o processo ganha rastreabilidade no Workflow.", "modulos/page.tsx:123")
add_field_row(t, "Pill 1",     "Menos retrabalho",                                           "modulos/page.tsx:127")
add_field_row(t, "Pill 2",     "Menos ruído de versão",                                      "modulos/page.tsx:128")
add_field_row(t, "Pill 3",     "Mais governança",                                            "modulos/page.tsx:129")
spacer(doc)

add_section_heading(doc, "Relatórios section")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",    "Relatórios",                                                 "modulos/page.tsx:136")
add_field_row(t, "Heading",    "O módulo de viabilidade entrega material para análise, não só telas de sistema.", "modulos/page.tsx:138")
add_field_row(t, "Subheading", "Os exemplos de relatórios da pasta de conteúdo reforçam que a decisão depende de premissas, fluxos, comparativos, tabelas e indicadores exportáveis.", "modulos/page.tsx:140")
add_field_row(t, "Report 1",   "Premissas e cenários — Registra condições comerciais, obra, financiamento, permutas e parâmetros que sustentam cada versão do estudo.", "modulos/page.tsx:157")
add_field_row(t, "Report 2",   "Fluxos de caixa — Permite leitura sintética ou analítica, nominal, indexada ou a valor presente, com visão do projeto e dos participantes.", "modulos/page.tsx:162")
add_field_row(t, "Report 3",   "Sensibilidade e indicadores — Mostra como TIR, VPL, margem, ROI, exposição de caixa e payback reagem a mudanças nas variáveis críticas.", "modulos/page.tsx:167")
spacer(doc)

add_section_heading(doc, "Exemplos de Relatório section")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",    "Exemplos de relatório",                                      "modulos/page.tsx:176")
add_field_row(t, "Heading",    "O VIABIL entrega documentos para análise, não só telas de sistema.", "modulos/page.tsx:178")
add_field_row(t, "Subheading", "Os relatórios saem do módulo de Viabilidade e do Acompanhamento prontos para comitê, sócios, investidores e parceiros. Os exemplos abaixo são amostras reais geradas pela plataforma.", "modulos/page.tsx:180")
spacer(doc)

add_section_heading(doc, "CTA Band")
t = make_section_table(doc)
add_field_row(t, "Heading",    "Quer entender quais módulos fazem sentido para sua operação?",  "modulos/page.tsx:330")
add_field_row(t, "Subheading", "A demonstração pode focar no estágio da sua empresa: originação, aprovação de novos projetos, acompanhamento ou visão consolidada.", "modulos/page.tsx:332")
spacer(doc)

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
#  7. SEGMENTOS PAGE
# ─────────────────────────────────────────────────────────────────────────────
add_page_heading(doc, "7. Segmentos Page", "app/(public)/segmentos/page.tsx")

add_section_heading(doc, "Page metadata")
t = make_section_table(doc)
add_field_row(t, "Page title",       "Segmentos | VIABIL",                                          "segmentos/page.tsx:5")
add_field_row(t, "Meta description", "O VIABIL atende incorporação residencial, casas, loteamentos e outros segmentos do mercado imobiliário com modelos parametrizáveis.", "segmentos/page.tsx:6")
spacer(doc)

add_section_heading(doc, "Hero")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",   "Segmentos",                                                   "segmentos/page.tsx:71")
add_field_row(t, "H1",        "Começa no residencial. Avança para todo o real estate.",      "segmentos/page.tsx:73")
add_field_row(t, "Subtitle",  "A prioridade do VIABIL é incorporação residencial, casas e loteamentos. A mesma base metodológica também atende renda, participações, originação, consultorias e proprietários de áreas.", "segmentos/page.tsx:75")
spacer(doc)

add_section_heading(doc, "Foco Principal — Incorporação residencial")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",    "Foco principal",                                             "segmentos/page.tsx:87")
add_field_row(t, "Heading",    "Incorporação, casas e loteamentos vêm primeiro.",            "segmentos/page.tsx:89")
add_field_row(t, "Subheading", "A estrutura segue a orientação do proprietário: apresentar os segmentos residenciais como o núcleo do posicionamento e tratar os demais como suportados.", "segmentos/page.tsx:91")
add_field_row(t, "Seg 1 name", "Incorporação residencial",                                   "segmentos/page.tsx seg 1")
add_field_row(t, "Seg 1 badge","Principal",                                                  "segmentos/page.tsx seg 1")
add_field_row(t, "Seg 1 challenge","Ciclos longos, terrenos caros, aprovações, funding, variação de custos e risco comercial tornam cada decisão de aquisição ou lançamento crítica.", "segmentos/page.tsx seg 1")
add_field_row(t, "Seg 1 value","O VIABIL organiza VGV, velocidade de vendas, permutas, custos, financiamento, indicadores e acompanhamento em uma visão que serve ao analista e ao comitê executivo.", "segmentos/page.tsx seg 1")
add_field_row(t, "Seg 1 detail 1","Aquisição de terreno",                                   "segmentos/page.tsx seg 1")
add_field_row(t, "Seg 1 detail 2","Lançamento e tabela de vendas",                          "segmentos/page.tsx seg 1")
add_field_row(t, "Seg 1 detail 3","Previsto x realizado",                                   "segmentos/page.tsx seg 1")
add_field_row(t, "Seg 1 detail 4","Consolidação de SPEs",                                   "segmentos/page.tsx seg 1")
spacer(doc)

add_section_heading(doc, "Foco Principal — Casas e condomínios horizontais")
t = make_section_table(doc)
add_field_row(t, "Seg 2 name",   "Casas e condomínios horizontais",                         "segmentos/page.tsx seg 2")
add_field_row(t, "Seg 2 badge",  "Residencial",                                             "segmentos/page.tsx seg 2")
add_field_row(t, "Seg 2 challenge","Produtos horizontais exigem controle de fases, infraestrutura, personalizações, absorção comercial e custos por tipologia.", "segmentos/page.tsx seg 2")
add_field_row(t, "Seg 2 value",  "A parametrização permite simular modelos de casas, sobrados, condomínios fechados e operações com diferentes ritmos de venda e entrega.", "segmentos/page.tsx seg 2")
add_field_row(t, "Seg 2 detail 1","Fases de lançamento",                                    "segmentos/page.tsx seg 2")
add_field_row(t, "Seg 2 detail 2","Custos por tipologia",                                   "segmentos/page.tsx seg 2")
add_field_row(t, "Seg 2 detail 3","Absorção de unidades",                                   "segmentos/page.tsx seg 2")
add_field_row(t, "Seg 2 detail 4","Estratégias de preço",                                   "segmentos/page.tsx seg 2")
spacer(doc)

add_section_heading(doc, "Foco Principal — Loteamentos e urbanização")
t = make_section_table(doc)
add_field_row(t, "Seg 3 name",   "Loteamentos e urbanização",                               "segmentos/page.tsx seg 3")
add_field_row(t, "Seg 3 badge",  "Residencial",                                             "segmentos/page.tsx seg 3")
add_field_row(t, "Seg 3 challenge","Licenças, infraestrutura, parceria com terrenistas e longos ciclos de capital próprio exigem visão financeira disciplinada.", "segmentos/page.tsx seg 3")
add_field_row(t, "Seg 3 value",  "O VIABIL contempla permutas, infraestrutura, tabelas simultâneas, financiamento direto, securitização e acompanhamento do plano aprovado.", "segmentos/page.tsx seg 3")
add_field_row(t, "Seg 3 detail 1","Permuta física e financeira",                            "segmentos/page.tsx seg 3")
add_field_row(t, "Seg 3 detail 2","Infraestrutura e terraplanagem",                         "segmentos/page.tsx seg 3")
add_field_row(t, "Seg 3 detail 3","Carteira e securitização",                               "segmentos/page.tsx seg 3")
add_field_row(t, "Seg 3 detail 4","Análise por fase",                                       "segmentos/page.tsx seg 3")
spacer(doc)

add_section_heading(doc, "Abrangência — Outros segmentos")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",    "Abrangência",                                               "segmentos/page.tsx:115")
add_field_row(t, "Heading",    "Outros segmentos sem roubar o foco.",                        "segmentos/page.tsx:117")
add_field_row(t, "Subheading", "Corporativo, logística, shopping, fundos, consultorias e proprietários de áreas continuam importantes. A página apenas não deve abrir com eles.", "segmentos/page.tsx:119")
add_field_row(t, "Seg 4 name", "Corporativo e locação",                                     "segmentos/page.tsx secondary seg 1")
add_field_row(t, "Seg 4 text", "Galpões logísticos, BTS, lajes, malls e ativos com renda, cap rate, vacância e contratos.", "segmentos/page.tsx secondary seg 1")
add_field_row(t, "Seg 5 name", "Investimentos e participações",                             "segmentos/page.tsx secondary seg 2")
add_field_row(t, "Seg 5 text", "Sócios, fundos, investidores, permutantes e estruturas com retorno individualizado.", "segmentos/page.tsx secondary seg 2")
add_field_row(t, "Seg 6 name", "Originação e desenvolvimento",                             "segmentos/page.tsx secondary seg 3")
add_field_row(t, "Seg 6 text", "Consultorias, áreas de novos negócios, imobiliárias e proprietários de áreas em análise.", "segmentos/page.tsx secondary seg 3")
spacer(doc)

add_section_heading(doc, "Parametrização")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",    "Parametrização",                                            "segmentos/page.tsx:135")
add_field_row(t, "Heading",    "A flexibilidade está nas premissas, não em promessas genéricas.", "segmentos/page.tsx:137")
add_field_row(t, "Subheading", "Cada segmento muda o peso das variáveis: preço, custo, infraestrutura, funding, permuta, aluguel, cap rate, velocidade de vendas, distrato, inadimplência e saída do investimento.", "segmentos/page.tsx:139")
spacer(doc)

add_section_heading(doc, "Método por Segmento")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",    "Método por segmento",                                       "segmentos/page.tsx:147")
add_field_row(t, "Heading",    "O desafio muda, mas a linguagem financeira permanece.",      "segmentos/page.tsx:149")
add_field_row(t, "Subheading", "A página agora deixa claro por que cada segmento exige premissas próprias, sem tirar o foco do residencial, casas e loteamentos.", "segmentos/page.tsx:151")
add_field_row(t, "Method 1 title","Incorporação residencial",                               "segmentos/page.tsx method 1")
add_field_row(t, "Method 1 text","O estudo combina aquisição do terreno, VGV, tabela de vendas, curva de obra, financiamento, SPEs e acompanhamento depois do lançamento.", "segmentos/page.tsx method 1")
add_field_row(t, "Method 2 title","Loteamentos",                                            "segmentos/page.tsx method 2")
add_field_row(t, "Method 2 text","A análise precisa suportar aprovação longa, infraestrutura pesada, parceria com terrenistas, carteira própria e fases de comercialização.", "segmentos/page.tsx method 2")
add_field_row(t, "Method 3 title","Ativos de renda",                                        "segmentos/page.tsx method 3")
add_field_row(t, "Method 3 text","Quando o retorno vem de locação ou saída de investimento, entram vacância, cap-rate, contratos, yield e estrutura de participação.", "segmentos/page.tsx method 3")
spacer(doc)

add_section_heading(doc, "Indicadores por Segmento section headers")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",    "Indicadores por segmento",                                  "segmentos/page.tsx:167")
add_field_row(t, "Heading",    "As premissas mudam, mas os indicadores e a linguagem permanecem.", "segmentos/page.tsx:169")
add_field_row(t, "Subheading", "Cada segmento tem variáveis específicas. O VIABIL parametriza e gera os mesmos indicadores — TIR, VPL, margem, exposição — adaptados à lógica de cada tipo de empreendimento.", "segmentos/page.tsx:171")
spacer(doc)

add_section_heading(doc, "CTA Band")
t = make_section_table(doc)
add_field_row(t, "Heading",    "Mostre seu tipo de empreendimento na demonstração.",         "segmentos/page.tsx:297")
add_field_row(t, "Subheading", "A conversa fica mais objetiva quando parte do segmento, do ciclo e das variáveis que realmente mudam a decisão.", "segmentos/page.tsx:299")
spacer(doc)

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
#  8. VERSÕES PAGE
# ─────────────────────────────────────────────────────────────────────────────
add_page_heading(doc, "8. Versões Page", "app/(public)/versoes/page.tsx")

add_section_heading(doc, "Page metadata")
t = make_section_table(doc)
add_field_row(t, "Page title",       "Versões | VIABIL",                                            "versoes/page.tsx:7")
add_field_row(t, "Meta description", "Conheça VIABIL Lite, Full e VIABIL Cloud. Versões compatíveis para troca de estudos entre parceiros.", "versoes/page.tsx:8")
spacer(doc)

add_section_heading(doc, "Hero")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",   "Versões",                                                     "versoes/page.tsx:87")
add_field_row(t, "H1",        "A mesma cultura VIABIL em estágios diferentes de operação.", "versoes/page.tsx:89")
add_field_row(t, "Subtitle",  "Lite, Full e Cloud mantêm a cultura VIABIL em diferentes formatos de operação. A escolha depende do tamanho da equipe, do volume de projetos e da necessidade de acesso, parametrização e acompanhamento.", "versoes/page.tsx:91")
spacer(doc)

add_section_heading(doc, "VIABIL Full — highlight panel")
t = make_section_table(doc)
add_field_row(t, "Badge",       "Versão corporativa",                                        "versoes/page.tsx Full panel")
add_field_row(t, "Heading",     "VIABIL é a plataforma completa para empresas que precisam padronizar decisões, parametrizar modelos e acompanhar o ciclo imobiliário com governança.", "versoes/page.tsx Full panel")
add_field_row(t, "Description", "Ideal para incorporadoras, loteadoras, construtoras e grupos com múltiplos projetos, equipes envolvidas e necessidade de customizações, integrações e visão consolidada.", "versoes/page.tsx Full panel")
add_field_row(t, "CTA",         "Solicitar proposta",                                        "versoes/page.tsx Full panel")
add_field_row(t, "Fact 1",      "Usuários ilimitados",                                       "versoes/page.tsx Full panel")
add_field_row(t, "Fact 2",      "Gestão de Terrenos, Viabilidade, Consolidação e Workflow",  "versoes/page.tsx Full panel")
add_field_row(t, "Fact 3",      "Parametrizações e relatórios personalizados",               "versoes/page.tsx Full panel")
add_field_row(t, "Fact 4",      "Base para acompanhamento planejado x realizado",            "versoes/page.tsx Full panel")
add_field_row(t, "Note",        "Para empresas que precisam transformar a metodologia VIABIL em padrão interno de decisão.", "versoes/page.tsx Full panel")
spacer(doc)

add_section_heading(doc, "VIABIL Lite — highlight panel")
t = make_section_table(doc)
add_field_row(t, "Badge",       "Versão acessível",                                          "versoes/page.tsx Lite panel")
add_field_row(t, "Heading",     "VIABIL Lite é a porta de entrada para a cultura VIABIL. A versão acessível para pequenas incorporadoras, loteadoras e consultorias que querem começar com o padrão do mercado.", "versoes/page.tsx Lite panel")
add_field_row(t, "Description", "Análise de viabilidade, cálculo de VGV, margens e simulações de cenários sem contrato de implantação. Ideal para quem está começando ou quer testar o método VIABIL antes de escalar.", "versoes/page.tsx Lite panel")
add_field_row(t, "CTA",         "Conhecer o VIABIL Lite",                                    "versoes/page.tsx Lite panel")
add_field_row(t, "Fact 1",      "Análise de viabilidade simplificada",                       "versoes/page.tsx Lite panel")
add_field_row(t, "Fact 2",      "Simulação de cenários básicos",                             "versoes/page.tsx Lite panel")
add_field_row(t, "Fact 3",      "Relatório de viabilidade em PDF",                           "versoes/page.tsx Lite panel")
add_field_row(t, "Fact 4",      "Sem necessidade de implantação",                            "versoes/page.tsx Lite panel")
add_field_row(t, "Note",        "Ideal para consultores, pequenas incorporadoras e profissionais em início de carreira.", "versoes/page.tsx Lite panel")
spacer(doc)

add_section_heading(doc, "Linha VIABIL — version cards")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",         "Linha VIABIL",                                          "versoes/page.tsx:151")
add_field_row(t, "Heading",         "Escolha pela maturidade da operação, não só pela lista de recursos.", "versoes/page.tsx:153")
add_field_row(t, "Subheading",      "A página foi ajustada para não prometer treinamento, pagamento, login ou recursos fora do escopo público do site.", "versoes/page.tsx:155")
add_field_row(t, "Lite tag",        "Porta de entrada",                                      "versoes/page.tsx card Lite")
add_field_row(t, "Lite profile",    "Pequenas empresas, consultores e desenvolvedores em início de estruturação.", "versoes/page.tsx card Lite")
add_field_row(t, "Lite description","Versão acessível para iniciar a cultura VIABIL com os principais recursos de viabilidade, até 2 licenças e sem customizações.", "versoes/page.tsx card Lite")
add_field_row(t, "Lite fact 1",     "Estudos compatíveis com Full",                          "versoes/page.tsx card Lite")
add_field_row(t, "Lite fact 2",     "Principais recursos de viabilidade",                    "versoes/page.tsx card Lite")
add_field_row(t, "Lite fact 3",     "Relatórios pré-formatados",                             "versoes/page.tsx card Lite")
add_field_row(t, "Lite fact 4",     "Investimento inicial mais leve",                        "versoes/page.tsx card Lite")
add_field_row(t, "Lite button",     "Solicitar proposta",                                    "versoes/page.tsx card Lite")
add_field_row(t, "Full tag",        "Padrão corporativo",                                    "versoes/page.tsx card Full")
add_field_row(t, "Full profile",    "Empresas com múltiplos projetos simultâneos e equipes envolvidas no processo.", "versoes/page.tsx card Full")
add_field_row(t, "Full description","Versão padrão para empresas em desenvolvimento, com usuários ilimitados, parametrizações, integrações e módulos de gestão do ciclo.", "versoes/page.tsx card Full")
add_field_row(t, "Full fact 1",     "Gestão de Terrenos",                                    "versoes/page.tsx card Full")
add_field_row(t, "Full fact 2",     "Viabilidade",                                           "versoes/page.tsx card Full")
add_field_row(t, "Full fact 3",     "Consolidação",                                          "versoes/page.tsx card Full")
add_field_row(t, "Full fact 4",     "Workflow de Tarefas",                                   "versoes/page.tsx card Full")
add_field_row(t, "Full button",     "Solicitar proposta",                                    "versoes/page.tsx card Full")
spacer(doc)

add_section_heading(doc, "VIABIL Cloud")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",     "VIABIL Cloud",                                              "versoes/page.tsx:193")
add_field_row(t, "Heading",     "Acesso em nuvem, com segurança e auditabilidade.",          "versoes/page.tsx:195")
add_field_row(t, "Description", "O VIABIL Cloud é o modelo por assinatura hospedado em infraestrutura Oracle Cloud + Sky.One. Permite acesso de qualquer dispositivo, sem investimento inicial de licença e com confidencialidade dos dados.", "versoes/page.tsx:197")
spacer(doc)

add_section_heading(doc, "Como Escolher")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",    "Como escolher",                                              "versoes/page.tsx:232")
add_field_row(t, "Heading",    "A versão certa depende do processo que a empresa precisa sustentar.", "versoes/page.tsx:234")
add_field_row(t, "Subheading", "A compatibilidade entre formatos mantém a troca de estudos com parceiros. O que muda é profundidade, parametrização, implantação, acesso em nuvem e acompanhamento.", "versoes/page.tsx:236")
add_field_row(t, "Choice 1 title","Quando Lite faz sentido",                                 "versoes/page.tsx choice 1")
add_field_row(t, "Choice 1 text","A empresa quer entrar na cultura VIABIL com investimento mais acessível, poucos usuários e foco nos principais estudos de viabilidade.", "versoes/page.tsx choice 1")
add_field_row(t, "Choice 2 title","Quando Full faz sentido",                                 "versoes/page.tsx choice 2")
add_field_row(t, "Choice 2 text","A operação já tem vários projetos, equipes envolvidas, necessidade de usuários ilimitados, parametrizações e troca de estudos com parceiros.", "versoes/page.tsx choice 2")
add_field_row(t, "Choice 3 title","Quando Acompanhamento faz sentido",                      "versoes/page.tsx choice 3")
add_field_row(t, "Choice 3 text","A maturidade exige comparar planejado, revisado e realizado, importar dados e replanejar ações durante a vida do empreendimento.", "versoes/page.tsx choice 3")
add_field_row(t, "Choice 4 title","Quando Cloud faz sentido",                               "versoes/page.tsx choice 4")
add_field_row(t, "Choice 4 text","O acesso em nuvem reduz investimento inicial de licença e simplifica uso remoto com segurança, confidencialidade e auditabilidade.", "versoes/page.tsx choice 4")
spacer(doc)

add_section_heading(doc, "Relatório do Investidor")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",    "Relatório do investidor",                                    "versoes/page.tsx:252")
add_field_row(t, "Heading",    "O VIABIL entrega visão detalhada por participante.",         "versoes/page.tsx:254")
add_field_row(t, "Subheading", "Sócios, investidores e permutantes têm fluxo e indicadores individualizados: capital investido, retorno, MTIR e cronograma de recebimentos. Exemplo real gerado pela plataforma.", "versoes/page.tsx:256")
add_field_row(t, "Button",     "Ver relatório completo ↗",                                   "versoes/page.tsx:259")
spacer(doc)

add_section_heading(doc, "CTA Band")
t = make_section_table(doc)
add_field_row(t, "Heading",    "Não sabe qual versão escolher?",                             "versoes/page.tsx:293")
add_field_row(t, "Subheading", "Uma conversa rápida ajuda a separar o que é essencial agora do que deve entrar em uma implantação mais completa.", "versoes/page.tsx:295")
spacer(doc)

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
#  9. SERVIÇOS PAGE
# ─────────────────────────────────────────────────────────────────────────────
add_page_heading(doc, "9. Serviços Page", "app/(public)/servicos/page.tsx")

add_section_heading(doc, "Page metadata")
t = make_section_table(doc)
add_field_row(t, "Page title",       "Serviços | VIABIL",                                           "servicos/page.tsx:5")
add_field_row(t, "Meta description", "Suporte ao usuário, assessoria operacional, implantação, parametrização, customizações e integrações para clientes VIABIL.", "servicos/page.tsx:6")
spacer(doc)

add_section_heading(doc, "Hero")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",  "Serviços",                                                     "servicos/page.tsx:54")
add_field_row(t, "H1",       "Tecnologia sem conteúdo é pouco eficaz.",                      "servicos/page.tsx:56")
add_field_row(t, "Subtitle", "A BDK Solutions opera full-service: desenvolvimento, suporte, treinamento, consultoria, implantação, customização e integração para que o VIABIL se encaixe na realidade de cada cliente.", "servicos/page.tsx:58")
spacer(doc)

add_section_heading(doc, "Full-service section")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",    "Full-service",                                               "servicos/page.tsx:70")
add_field_row(t, "Heading",    "Serviço faz parte do valor do VIABIL.",                      "servicos/page.tsx:72")
add_field_row(t, "Subheading", "O software carrega metodologia, mas a adoção real acontece quando suporte, parametrização e consultoria ajudam a empresa a transformar processo.", "servicos/page.tsx:74")
spacer(doc)

add_section_heading(doc, "Serviço 1 — Suporte ao Usuário")
t = make_section_table(doc)
add_field_row(t, "Name",       "Suporte ao Usuário",                                         "servicos/page.tsx svc 1")
add_field_row(t, "Tag",        "Dia a dia",                                                  "servicos/page.tsx svc 1")
add_field_row(t, "Description","Atendimento para dúvidas operacionais, técnicas e conceituais. A equipe entende viabilidade financeira imobiliária, não apenas funcionamento de software.", "servicos/page.tsx svc 1")
add_field_row(t, "Fact 1",     "300+ atendimentos semanais",                                 "servicos/page.tsx svc 1")
add_field_row(t, "Fact 2",     "Dúvidas de uso e conceito",                                  "servicos/page.tsx svc 1")
add_field_row(t, "Fact 3",     "Atualizações e novos recursos",                              "servicos/page.tsx svc 1")
spacer(doc)

add_section_heading(doc, "Serviço 2 — Assessoria Operacional")
t = make_section_table(doc)
add_field_row(t, "Name",       "Assessoria Operacional",                                     "servicos/page.tsx svc 2")
add_field_row(t, "Tag",        "Casos reais",                                                "servicos/page.tsx svc 2")
add_field_row(t, "Description","Especialistas VIABIL trabalham junto com a equipe do cliente em estudos, validação de premissas e discussão de boas práticas de mercado.", "servicos/page.tsx svc 2")
add_field_row(t, "Fact 1",     "Realização conjunta de estudos",                             "servicos/page.tsx svc 2")
add_field_row(t, "Fact 2",     "Validação de critérios",                                     "servicos/page.tsx svc 2")
add_field_row(t, "Fact 3",     "Troca de melhores práticas",                                 "servicos/page.tsx svc 2")
spacer(doc)

add_section_heading(doc, "Serviço 3 — Implantação e Parametrização")
t = make_section_table(doc)
add_field_row(t, "Name",       "Implantação e Parametrização",                               "servicos/page.tsx svc 3")
add_field_row(t, "Tag",        "Adoção",                                                     "servicos/page.tsx svc 3")
add_field_row(t, "Description","Configuração inicial de premissas, indicadores, curvas, plano de contas, estudos-modelo e homologação para que a empresa adote o padrão VIABIL.", "servicos/page.tsx svc 3")
add_field_row(t, "Fact 1",     "Modelo de importação",                                       "servicos/page.tsx svc 3")
add_field_row(t, "Fact 2",     "Geração de conteúdo",                                        "servicos/page.tsx svc 3")
add_field_row(t, "Fact 3",     "Testes e homologação",                                       "servicos/page.tsx svc 3")
spacer(doc)

add_section_heading(doc, "Serviço 4 — Customizações")
t = make_section_table(doc)
add_field_row(t, "Name",       "Customizações",                                              "servicos/page.tsx svc 4")
add_field_row(t, "Tag",        "Sob medida",                                                 "servicos/page.tsx svc 4")
add_field_row(t, "Description","Relatórios, indicadores e extensões funcionais para empresas com demandas específicas, sempre desenvolvidas pela equipe dedicada ao VIABIL.", "servicos/page.tsx svc 4")
add_field_row(t, "Fact 1",     "80+ projetos realizados",                                    "servicos/page.tsx svc 4")
add_field_row(t, "Fact 2",     "Indicadores personalizados",                                 "servicos/page.tsx svc 4")
add_field_row(t, "Fact 3",     "Relatórios executivos",                                      "servicos/page.tsx svc 4")
spacer(doc)

add_section_heading(doc, "Serviço 5 — Integrações")
t = make_section_table(doc)
add_field_row(t, "Name",       "Integrações",                                                "servicos/page.tsx svc 5")
add_field_row(t, "Tag",        "Dados",                                                      "servicos/page.tsx svc 5")
add_field_row(t, "Description","Caminhos de integração com os principais ERPs do mercado imobiliário, conforme escopo do projeto, por layouts, exportações ou acesso estruturado a dados.", "servicos/page.tsx svc 5")
add_field_row(t, "Fact 1",     "Layouts de importação",                                      "servicos/page.tsx svc 5")
add_field_row(t, "Fact 2",     "De-para de plano de contas",                                 "servicos/page.tsx svc 5")
add_field_row(t, "Fact 3",     "Base para acompanhamento",                                   "servicos/page.tsx svc 5")
spacer(doc)

add_section_heading(doc, "Conhecimento Aplicado")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",    "Conhecimento aplicado",                                      "servicos/page.tsx:98")
add_field_row(t, "Heading",    "Suporte, implantação e treinamento falam a língua do mercado.", "servicos/page.tsx:100")
add_field_row(t, "Subheading", "Os consultores VIABIL ajudam em dúvidas conceituais, critérios de estudo, parametrização e práticas que surgem em incorporações, casas, loteamentos e demais segmentos.", "servicos/page.tsx:102")
add_field_row(t, "Metric 1",   "300+ atendimentos por semana",                               "servicos/page.tsx:106")
add_field_row(t, "Metric 2",   "120+ treinamentos por ano",                                  "servicos/page.tsx:110")
spacer(doc)

add_section_heading(doc, "Implantação")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",    "Implantação",                                                "servicos/page.tsx:120")
add_field_row(t, "Heading",    "Parametrizar é traduzir a forma de trabalho da empresa para o VIABIL.", "servicos/page.tsx:122")
add_field_row(t, "Subheading", "A implantação foi detalhada para explicar o caminho entre contratar o sistema e usar a metodologia com segurança em estudos reais.", "servicos/page.tsx:124")
add_field_row(t, "Step 1",     "1. Modelo de importação — Definição de layouts, plano de contas e de-para para que dados externos entrem com consistência.", "servicos/page.tsx impl 1")
add_field_row(t, "Step 2",     "2. Conteúdo e parâmetros — Configuração de premissas, curvas, indicadores, estudos-modelo e relatórios alinhados à realidade do cliente.", "servicos/page.tsx impl 2")
add_field_row(t, "Step 3",     "3. Testes e homologação — Rodadas de validação com casos reais antes de consolidar o uso pela equipe.", "servicos/page.tsx impl 3")
add_field_row(t, "Step 4",     "4. Adoção assistida — Apoio inicial para dúvidas operacionais, conceituais e ajustes finos após a implantação.", "servicos/page.tsx impl 4")
spacer(doc)

add_section_heading(doc, "Tabela de Vendas section")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",    "Tabela de vendas",                                           "servicos/page.tsx:139")
add_field_row(t, "Heading",    "A parametrização define o padrão de análise da empresa.",    "servicos/page.tsx:141")
add_field_row(t, "Subheading", "Na implantação, o VIABIL é configurado com as premissas, plano de contas, curvas, indicadores e tabelas de venda do cliente. O resultado é um modelo-padrão alinhado à realidade da operação — não um template genérico.", "servicos/page.tsx:143")
add_field_row(t, "Link text",  "Ver exemplo de tabela real",                                 "servicos/page.tsx:190")
spacer(doc)

add_section_heading(doc, "CTA Band")
t = make_section_table(doc)
add_field_row(t, "Heading",    "Precisa parametrizar o VIABIL para sua operação?",           "servicos/page.tsx:195")
add_field_row(t, "Subheading", "A conversa comercial pode mapear versão, implantação, customizações e integrações sem prometer escopo antes da análise técnica.", "servicos/page.tsx:197")
spacer(doc)

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
#  10. SOBRE PAGE
# ─────────────────────────────────────────────────────────────────────────────
add_page_heading(doc, "10. Sobre Page", "app/(public)/sobre/page.tsx")

add_section_heading(doc, "Page metadata")
t = make_section_table(doc)
add_field_row(t, "Page title",       "Sobre | VIABIL e BDK Solutions",                              "sobre/page.tsx:6")
add_field_row(t, "Meta description", "Conheça a BDK Solutions, empresa por trás do VIABIL, fundada em 1995 por Eli Wolf.", "sobre/page.tsx:7")
spacer(doc)

add_section_heading(doc, "Hero")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",  "Sobre",                                                        "sobre/page.tsx:45")
add_field_row(t, "H1",       "A empresa por trás do padrão VIABIL.",                         "sobre/page.tsx:47")
add_field_row(t, "Subtitle", "A BDK Solutions desenvolve, comercializa, treina, implanta e suporta o VIABIL em todo o Brasil. O produto é a identidade pública; a BDK é a estrutura que sustenta conhecimento, serviço e evolução.", "sobre/page.tsx:49")
spacer(doc)

add_section_heading(doc, "Manifesto")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",    "Manifesto",                                                  "sobre/page.tsx:62")
add_field_row(t, "Heading",    "Nascemos dentro do universo da incorporação imobiliária.",   "sobre/page.tsx:64")
add_field_row(t, "Subheading", "A BDK absorve demandas e práticas de mercado e traduz esse conhecimento em conceitos, recursos, tecnologia e serviços. É por isso que o VIABIL não soa como um sistema genérico adaptado ao setor.", "sobre/page.tsx:66")
add_field_row(t, "Timeline 1995","Fundação da BDK Solutions, com atuação voltada a conhecimento, tecnologia e negócios imobiliários.", "sobre/page.tsx timeline 1")
add_field_row(t, "Timeline Mercado","O VIABIL nasce dentro do universo da incorporação, absorvendo demandas reais de clientes e práticas do setor.", "sobre/page.tsx timeline 2")
add_field_row(t, "Timeline Hoje","A plataforma é usada por 600+ empresas e segue evoluindo com suporte, consultoria, treinamento e desenvolvimento dedicado.", "sobre/page.tsx timeline 3")
spacer(doc)

add_section_heading(doc, "Liderança")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",  "Liderança",                                                    "sobre/page.tsx:77")
add_field_row(t, "Name",     "Eli Wolf",                                                     "sobre/page.tsx:79")
add_field_row(t, "Bio",      "Diretor Executivo e idealizador do VIABIL. Une formação em Tecnologia e Administração de Negócios a mais de 30 anos no mercado imobiliário, participando de produto, treinamento, consultoria e eventos do setor.", "sobre/page.tsx:81")
spacer(doc)

add_section_heading(doc, "Propósito")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",    "Propósito",                                                  "sobre/page.tsx:100")
add_field_row(t, "Heading",    "Democratizar e padronizar melhores práticas.",               "sobre/page.tsx:102")
add_field_row(t, "Subheading", "A missão é contribuir com o desenvolvimento do setor por meio de softwares e serviços que levam metodologia, segurança e eficiência para decisões imobiliárias.", "sobre/page.tsx:104")
add_field_row(t, "Metric 1",   "600+ empresas implementadas",                               "sobre/page.tsx:108")
add_field_row(t, "Metric 2",   "8.000+ profissionais treinados",                            "sobre/page.tsx:112")
spacer(doc)

add_section_heading(doc, "Manifesto Aplicado")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",    "Manifesto aplicado",                                         "sobre/page.tsx:122")
add_field_row(t, "Heading",    "O DNA imobiliário aparece no produto e no serviço.",         "sobre/page.tsx:124")
add_field_row(t, "Subheading", "A BDK não trata o VIABIL como software genérico. A empresa absorve práticas do mercado e devolve isso em metodologia, suporte e evolução contínua.", "sobre/page.tsx:126")
add_field_row(t, "Item 1 title","Tecnologia com conteúdo",                                   "sobre/page.tsx manifesto 1")
add_field_row(t, "Item 1 text","O VIABIL traduz prática de mercado em recursos, premissas, relatórios e serviços. Sem conteúdo, tecnologia perde eficácia.", "sobre/page.tsx manifesto 1")
add_field_row(t, "Item 2 title","Clientes como parceiros",                                   "sobre/page.tsx manifesto 2")
add_field_row(t, "Item 2 text","A evolução do produto nasce da troca contínua com incorporadoras, loteadoras, consultorias, investidores e equipes usuárias.", "sobre/page.tsx manifesto 2")
add_field_row(t, "Item 3 title","Responsabilidade",                                          "sobre/page.tsx manifesto 3")
add_field_row(t, "Item 3 text","Ser referência em viabilidade financeira exige cuidado com cálculo, método, treinamento e suporte em decisões de alto impacto.", "sobre/page.tsx manifesto 3")
spacer(doc)

add_section_heading(doc, "Valores")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",    "Valores",                                                    "sobre/page.tsx:141")
add_field_row(t, "Heading",    "O que sustenta a relação com o mercado.",                    "sobre/page.tsx:143")
add_field_row(t, "Subheading", "Os princípios que guiam como a BDK Solutions atua, evolui e se relaciona com clientes e o mercado.", "sobre/page.tsx:145")
add_field_row(t, "Value 1",    "Dedicação ao cliente — Proximidade para entender a realidade da empresa e apoiar o uso correto da metodologia.", "sobre/page.tsx value 1")
add_field_row(t, "Value 2",    "Orientação à solução — Foco em resolver problemas concretos de decisão, operação e governança.", "sobre/page.tsx value 2")
add_field_row(t, "Value 3",    "Responsabilidade — Cálculos, premissas e suporte tratados com o peso de decisões financeiras críticas.", "sobre/page.tsx value 3")
add_field_row(t, "Value 4",    "Conhecimento — Tecnologia acompanhada de conteúdo prático do setor imobiliário.", "sobre/page.tsx value 4")
add_field_row(t, "Value 5",    "Excelência — Melhoria contínua no software, nos serviços e nas práticas internas.", "sobre/page.tsx value 5")
add_field_row(t, "Value 6",    "Crescimento sustentável — Decisões melhores para empresas, empreendimentos e para a economia do setor.", "sobre/page.tsx value 6")
spacer(doc)

add_section_heading(doc, "CTA Band")
t = make_section_table(doc)
add_field_row(t, "Heading",    "Conheça a metodologia por trás do VIABIL.",                  "sobre/page.tsx:157")
add_field_row(t, "Subheading", "A plataforma é produto de décadas de prática no mercado imobiliário, suporte próximo e evolução contínua com clientes parceiros.", "sobre/page.tsx:159")
spacer(doc)

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
#  11. CONTATO PAGE
# ─────────────────────────────────────────────────────────────────────────────
add_page_heading(doc, "11. Contato Page", "app/(public)/contato/page.tsx")

add_section_heading(doc, "Page metadata")
t = make_section_table(doc)
add_field_row(t, "Page title",       "Contato | VIABIL",                                            "contato/page.tsx:6")
add_field_row(t, "Meta description", "Fale com o time VIABIL para solicitar demonstração, proposta comercial ou conversar sobre parcerias.", "contato/page.tsx:7")
spacer(doc)

add_section_heading(doc, "Hero")
t = make_section_table(doc)
add_field_row(t, "Eyebrow",  "Contato",                                                      "contato/page.tsx:33")
add_field_row(t, "H1",       "Fale com um especialista VIABIL.",                             "contato/page.tsx:35")
add_field_row(t, "Subtitle", "Solicite uma demonstração, uma proposta ou uma conversa sobre a melhor forma de estruturar viabilidade, acompanhamento e governança financeira na sua operação imobiliária.", "contato/page.tsx:37")
spacer(doc)

add_section_heading(doc, "Contact paths")
t = make_section_table(doc)
add_field_row(t, "Section eyebrow",  "Caminhos",                                             "contato/page.tsx:50")
add_field_row(t, "Section heading",  "Escolha o assunto e avance pelo WhatsApp.",            "contato/page.tsx:52")
add_field_row(t, "Section subheading","O WhatsApp é o canal principal de conversão do site. O formulário ao lado monta a mensagem com os dados mínimos para contato.", "contato/page.tsx:54")
add_field_row(t, "Path 1 title",     "Agende uma apresentação",                              "contato/page.tsx path 1")
add_field_row(t, "Path 1 desc",      "Para ver o VIABIL aplicado ao seu segmento, com foco nas decisões que sua equipe precisa sustentar.", "contato/page.tsx path 1")
add_field_row(t, "Path 1 button",    "Iniciar conversa",                                     "contato/page.tsx path 1")
add_field_row(t, "Path 2 title",     "Solicite uma proposta",                                "contato/page.tsx path 2")
add_field_row(t, "Path 2 desc",      "Para discutir versão, implantação, parametrização e serviços adequados ao estágio da empresa.", "contato/page.tsx path 2")
add_field_row(t, "Path 2 button",    "Iniciar conversa",                                     "contato/page.tsx path 2")
add_field_row(t, "Path 3 title",     "Converse sobre parcerias",                             "contato/page.tsx path 3")
add_field_row(t, "Path 3 desc",      "Para consultorias, parceiros, instituições e iniciativas ligadas ao mercado imobiliário.", "contato/page.tsx path 3")
add_field_row(t, "Path 3 button",    "Iniciar conversa",                                     "contato/page.tsx path 3")
spacer(doc)

add_section_heading(doc, "Contact Form")
t = make_section_table(doc)
add_field_row(t, "Form eyebrow",  "Formulário",                                              "contato/page.tsx:79")
add_field_row(t, "Form heading",  "Envie os dados principais.",                              "contato/page.tsx:81")
add_field_row(t, "Form subheading","Os dados são registrados para que a equipe VIABIL retorne com o contexto correto.", "contato/page.tsx:83")
spacer(doc)

add_section_heading(doc, "Contact Form — Fields (ContactWhatsAppForm.tsx)")
t = make_section_table(doc)
add_field_row(t, "Field: Nome label",      "Nome",                                          "ContactWhatsAppForm.tsx — field Nome")
add_field_row(t, "Field: Empresa label",   "Empresa",                                       "ContactWhatsAppForm.tsx — field Empresa")
add_field_row(t, "Field: Cargo label",     "Cargo",                                         "ContactWhatsAppForm.tsx — field Cargo")
add_field_row(t, "Field: E-mail label",    "E-mail",                                        "ContactWhatsAppForm.tsx — field Email")
add_field_row(t, "Field: Telefone label",  "Telefone / WhatsApp",                           "ContactWhatsAppForm.tsx — field Tel")
add_field_row(t, "Field: Segmento label",  "Segmento",                                      "ContactWhatsAppForm.tsx — field Segmento")
add_field_row(t, "Segmento opt 1",         "Incorporação/Construção",                        "ContactWhatsAppForm.tsx — segmento options")
add_field_row(t, "Segmento opt 2",         "Loteamentos",                                   "ContactWhatsAppForm.tsx — segmento options")
add_field_row(t, "Segmento opt 3",         "Condomínios",                                   "ContactWhatsAppForm.tsx — segmento options")
add_field_row(t, "Segmento opt 4",         "Logísticos",                                    "ContactWhatsAppForm.tsx — segmento options")
add_field_row(t, "Segmento opt 5",         "Shopping",                                      "ContactWhatsAppForm.tsx — segmento options")
add_field_row(t, "Segmento opt 6",         "Corporativo / BTS",                             "ContactWhatsAppForm.tsx — segmento options")
add_field_row(t, "Segmento opt 7",         "Desenvolvimento Imobiliário / Originação",       "ContactWhatsAppForm.tsx — segmento options")
add_field_row(t, "Segmento opt 8",         "Consultoria",                                   "ContactWhatsAppForm.tsx — segmento options")
add_field_row(t, "Segmento opt 9",         "Proprietário de Área",                          "ContactWhatsAppForm.tsx — segmento options")
add_field_row(t, "Segmento opt 10",        "Participações / Investimentos",                  "ContactWhatsAppForm.tsx — segmento options")
add_field_row(t, "Field: Como conheceu",   "Como conheceu o VIABIL",                        "ContactWhatsAppForm.tsx — field ComoConheceu")
add_field_row(t, "Conheceu opt 1",         "Indicação",                                     "ContactWhatsAppForm.tsx — conheceu options")
add_field_row(t, "Conheceu opt 2",         "Eventos",                                       "ContactWhatsAppForm.tsx — conheceu options")
add_field_row(t, "Conheceu opt 3",         "Email marketing",                               "ContactWhatsAppForm.tsx — conheceu options")
add_field_row(t, "Conheceu opt 4",         "Redes sociais",                                 "ContactWhatsAppForm.tsx — conheceu options")
add_field_row(t, "Conheceu opt 5",         "Internet",                                      "ContactWhatsAppForm.tsx — conheceu options")
add_field_row(t, "Conheceu opt 6",         "Já usei VIABIL",                                "ContactWhatsAppForm.tsx — conheceu options")
add_field_row(t, "Field: Mensagem label",  "Mensagem",                                      "ContactWhatsAppForm.tsx — field Mensagem")
add_field_row(t, "Mensagem placeholder",   "Conte rapidamente o que sua equipe precisa analisar.", "ContactWhatsAppForm.tsx — mensagem placeholder")
add_field_row(t, "Submit button",          "Enviar dados",                                  "ContactWhatsAppForm.tsx — submit button")
add_field_row(t, "Submit loading",         "Enviando...",                                   "ContactWhatsAppForm.tsx — loading state")
add_field_row(t, "Validation error",       "Preencha nome, e-mail e telefone para continuar.", "ContactWhatsAppForm.tsx:41")
add_field_row(t, "Success message",        "Recebemos seus dados. Um especialista VIABIL entrará em contato.", "ContactWhatsAppForm.tsx — success")
add_field_row(t, "Error message",          "Não foi possível enviar agora. Tente novamente.", "ContactWhatsAppForm.tsx — error")
spacer(doc)

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
#  12. NEWSLETTER SIGNUP
# ─────────────────────────────────────────────────────────────────────────────
add_page_heading(doc, "12. Newsletter Signup", "components/marketing/NewsletterSignup.tsx")
t = make_section_table(doc)
add_field_row(t, "Section label",   "Viabilidade em Pauta",                                  "NewsletterSignup.tsx:49")
add_field_row(t, "H2",             "Receba conteúdos e atualizações do VIABIL.",             "NewsletterSignup.tsx:52")
add_field_row(t, "Description",    "Informações sobre viabilidade, mercado imobiliário, produto e temas que impactam decisões financeiras de incorporação.", "NewsletterSignup.tsx:54")
add_field_row(t, "Email label",    "E-mail profissional",                                    "NewsletterSignup.tsx:60")
add_field_row(t, "Name placeholder","Nome",                                                  "NewsletterSignup.tsx:68")
add_field_row(t, "Company placeholder","Empresa",                                            "NewsletterSignup.tsx:69")
add_field_row(t, "Subscribe button","Assinar",                                               "NewsletterSignup.tsx:64")
add_field_row(t, "Loading button",  "Enviando...",                                           "NewsletterSignup.tsx — loading")
add_field_row(t, "Validation error","Informe seu e-mail profissional.",                      "NewsletterSignup.tsx:17")
add_field_row(t, "Generic error",  "Não foi possível registrar sua inscrição agora.",        "NewsletterSignup.tsx:39")
add_field_row(t, "Success",        "Inscrição registrada. Obrigado por acompanhar o VIABIL.","NewsletterSignup.tsx:44")
spacer(doc)

# ─────────────────────────────────────────────────────────────────────────────
#  13. SKIP LINK / ACCESSIBILITY
# ─────────────────────────────────────────────────────────────────────────────
add_page_heading(doc, "13. Accessibility & Skip Links", "app/(public)/layout.tsx")
t = make_section_table(doc)
add_field_row(t, "Skip link",      "Ir para o conteúdo",                                     "layout.tsx:9")
add_field_row(t, "WhatsApp btn label","Falar com a equipe VIABIL pelo WhatsApp",             "WhatsAppFloatingButton.tsx:16")
add_field_row(t, "WhatsApp btn title","Falar pelo WhatsApp",                                 "WhatsAppFloatingButton.tsx:17")
spacer(doc)

# ─────────────────────────────────────────────────────────────────────────────
#  Save
# ─────────────────────────────────────────────────────────────────────────────
output_path = r"C:\GitHub\viabil_website\VIABIL_Website_Content_Review.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
